import tkinter as tk
from tkinter import ttk, messagebox
import sys
import os
import json
from datetime import datetime

# --- SAFE IMPORT FOR DOCX ---
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    root = tk.Tk()
    messagebox.showerror("Fehler", "Bibliothek 'python-docx' fehlt.\nBitte installieren: pip install python-docx")
    sys.exit(1)

def get_application_path():
    """Returns the base path where the application and config files are located."""
    if getattr(sys, 'frozen', False):
        application_path = os.path.dirname(sys.executable)
        if "Contents/MacOS" in application_path:
            application_path = os.path.abspath(os.path.join(application_path, "../../.."))
    else:
        application_path = os.path.dirname(os.path.abspath(__file__))
    return application_path

def get_template_path():
    return os.path.join(get_application_path(), "template.docx")

def get_config_path():
    return os.path.join(get_application_path(), "module_mapping.json")

def load_module_mapping():
    """Load module mapping from JSON file, create default if not exists."""
    config_path = get_config_path()
    
    # Default mapping (fallback)
    default_mapping = {
        "Elements of Mathematics": {"id": "MA4DSC1001", "name": "MA4DSC1001 Elements of Mathematics (WP)"},
        "Elements of Computer Science (Part 1)": {"id": "MA4DSC1002", "name": "MA4DSC1002 Elements of Computer Science"},
        "Elements of Computer Science (Part 2)": {"id": "MA4DSC1002", "name": "MA4DSC1002 Elements of Computer Science"},
        "Elements of Statistics": {"id": "MA4DSC1003", "name": "MA4DSC1003 Elements of Statistics"},
        "Statistical Programming with R": {"id": "MA4DSC1004", "name": "MA4DSC1004 Statistical Programming with R"},
        "Data Mining": {"id": "BA4WIN6008", "name": "BA4WIN6008 Data Mining"},
        "Big Data Analytics": {"id": "MA4DSC1008", "name": "MA4DSC1008 Big Data Analytics"},
        "Information Visualization": {"id": "MA4INF3022", "name": "MA4INF3022 Informationsvisualisierung"},
        "Semantic Technologies": {"id": "MA4WIN6010", "name": "MA4WIN6010 Semantische Technologien"},
        "Numerical Optimization for Data Science": {"id": "MA4FWB4503", "name": "MA4FWB4503 Ausgewählte Kapitel der Mathematik C"},
        "Fundamentals of Environmental Remote Sensing": {"id": "MA6FWB4302", "name": "MA6FWB4302 Fundamentals of Environmental Remote Sensing"},
        "Introduction to Geoinformatics": {"id": "MA6FWB4301", "name": "MA6FWB4301 Introduction to Geoinformatics"},
        "Deutsch: Grundkurs I (A1.1)": {"id": "MA2FWB1155", "name": "MA2FWB1155 Deutsch: Grundkurs I (A1.1)"},
        "Deutsch: Grundkurs II (A1.2)": {"id": "MA2FWB1156", "name": "MA2FWB1156 Deutsch: Grundkurs II (A1.2)"},
        "Deutsch: Aufbaukurs I (A2.1)": {"id": "MA2FWB1157", "name": "MA2FWB1157 Deutsch: Aufbaukurs I (A2.1)"},
        "Deutsch: Aufbaukurs II (A2.2)": {"id": "MA2FWB1158", "name": "MA2FWB1158 Deutsch: Aufbaukurs II (A2.2)"}
    }
    
    # If file doesn't exist, create it with default mapping
    if not os.path.exists(config_path):
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(default_mapping, f, indent=2, ensure_ascii=False)
            print(f"Created default config file: {config_path}")
        except Exception as e:
            messagebox.showwarning("Config Warning", 
                f"Could not create config file: {e}\nUsing default mapping.")
            return default_mapping
    
    # Load from file
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            mapping = json.load(f)
            print(f"Loaded {len(mapping)} modules from config file")
            return mapping
    except Exception as e:
        messagebox.showerror("Config Error", 
            f"Error loading module_mapping.json: {e}\nUsing default mapping.")
        return default_mapping

class RecognitionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Anerkennung NLP - Generator (Editierbar)")
        self.root.geometry("1200x900")

        # --- LOAD MODULE MAPPING FROM JSON ---
        self.module_map = load_module_mapping()
        
        self.rows_data = [] 

        # --- GUI SETUP ---
        info_frame = ttk.LabelFrame(root, text="Student Details")
        info_frame.pack(fill="x", padx=15, pady=10)

        # Zeile 0
        ttk.Label(info_frame, text="Anrede:").grid(row=0, column=0, sticky="e")
        self.gender_var = tk.StringVar()
        self.gender_combo = ttk.Combobox(info_frame, textvariable=self.gender_var, values=["Herr", "Frau"], state="readonly", width=10)
        self.gender_combo.current(0)
        self.gender_combo.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(info_frame, text="Name:*").grid(row=0, column=2, sticky="e")
        self.name_entry = ttk.Entry(info_frame, width=30)
        self.name_entry.grid(row=0, column=3, padx=5, pady=5)

        # Zeile 1
        ttk.Label(info_frame, text="Matrikelnr:*").grid(row=1, column=0, sticky="e")
        self.mat_entry = ttk.Entry(info_frame, width=15)
        self.mat_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(info_frame, text="Vorheriges Studium:*").grid(row=1, column=2, sticky="e")
        self.prev_study_entry = ttk.Entry(info_frame, width=35)
        self.prev_study_entry.insert(0, "M.Sc. Data Science (PO 2016/2021)")
        self.prev_study_entry.grid(row=1, column=3, padx=5, pady=5)

        # Zeile 2
        ttk.Label(info_frame, text="Einstufung (FS):*").grid(row=2, column=0, sticky="e")
        self.fs_entry = ttk.Entry(info_frame, width=15)
        self.fs_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(info_frame, text="Ziel-Studiengang:*").grid(row=2, column=2, sticky="e")
        self.target_po_entry = ttk.Entry(info_frame, width=35)
        self.target_po_entry.insert(0, "M.Sc. NLP (Version 2025)")
        self.target_po_entry.grid(row=2, column=3, padx=5, pady=5)

        # Input Frame
        input_frame = ttk.LabelFrame(root, text="Kurse eingeben")
        input_frame.pack(fill="x", padx=15, pady=5)
        
        # Headers
        ttk.Label(input_frame, text="1. Kurs auswählen", font=('bold')).grid(row=0, column=0, sticky="w", padx=5)
        ttk.Label(input_frame, text="2. Note (z.B. 2)", font=('bold')).grid(row=0, column=1, sticky="w", padx=5)
        ttk.Label(input_frame, text="Ziel-Modul (Auto)", font=('bold')).grid(row=0, column=2, sticky="w", padx=5)

        # 1. Course Input
        self.course_var = tk.StringVar()
        self.course_input = ttk.Combobox(input_frame, textvariable=self.course_var, width=35)
        self.course_input.grid(row=1, column=0, padx=5, pady=5)
        self.course_input.bind("<<ComboboxSelected>>", self.on_course_select)
        self.course_input.bind("<Return>", lambda e: self.grade_input.focus_set())
        
        # 2. Grade Input
        self.grade_input = ttk.Entry(input_frame, width=10)
        self.grade_input.grid(row=1, column=1, padx=5, pady=5)
        self.grade_input.bind("<Return>", self.add_row_event)

        # 3. Target Display (Readonly)
        self.target_var = tk.StringVar()
        self.target_display = ttk.Entry(input_frame, textvariable=self.target_var, width=35, state="readonly")
        self.target_display.grid(row=1, column=2, padx=5, pady=5)

        # Buttons Input
        ttk.Button(input_frame, text="Hinzufügen ⏎", command=self.add_row).grid(row=1, column=3, padx=10)
        ttk.Button(input_frame, text="Alles Leeren", command=self.clear_all).grid(row=1, column=4, padx=5)
        ttk.Button(input_frame, text="🔄 Config Neu Laden", command=self.reload_config).grid(row=1, column=5, padx=5)

        # --- TREEVIEW & EDIT CONTROLS ---
        
        tree_container = ttk.Frame(root)
        tree_container.pack(fill="both", expand=True, padx=15, pady=10)

        # Toolbar Frame für Bearbeiten/Löschen
        toolbar = ttk.Frame(tree_container)
        toolbar.pack(fill="x", pady=(0, 5))
        
        self.btn_edit = ttk.Button(toolbar, text="✏️ Auswahl bearbeiten", command=self.edit_selected, state="disabled")
        self.btn_edit.pack(side="left", padx=5)
        
        self.btn_del = ttk.Button(toolbar, text="🗑️ Auswahl löschen", command=self.delete_selected, state="disabled")
        self.btn_del.pack(side="left", padx=5)
        
        ttk.Label(toolbar, text="(Tipp: Doppelklick zum Bearbeiten, 'Entf' zum Löschen)", font=("Arial", 9, "italic"), foreground="gray").pack(side="left", padx=15)

        # Treeview
        self.tree = ttk.Treeview(tree_container, columns=("src", "grade", "tgt", "status"), show="headings", height=10)
        self.tree.heading("src", text="Gewählter Kurs")
        self.tree.heading("grade", text="Note")
        self.tree.heading("tgt", text="Ziel-Modul")
        self.tree.heading("status", text="Status")
        
        self.tree.column("src", width=250)
        self.tree.column("grade", width=80, anchor="center")
        self.tree.column("tgt", width=300)
        self.tree.column("status", width=150)
        
        self.tree.pack(fill="both", expand=True)
        self.tree.tag_configure('excluded', foreground='gray')
        
        # BINDINGS für Interaktion
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        self.tree.bind("<Double-1>", self.edit_selected)
        self.tree.bind("<Delete>", self.delete_selected)

        # Generate Button
        self.gen_btn = ttk.Button(root, text="Word Dokument erstellen", command=self.generate_document, state="disabled")
        self.gen_btn.pack(pady=15, ipadx=10, ipady=5)

        # Status bar with config info
        status_frame = ttk.Frame(root)
        status_frame.pack(fill="x", padx=15, pady=(0, 5))
        config_path = get_config_path()
        ttk.Label(status_frame, text=f"📁 Config: {config_path} ({len(self.module_map)} Module)", 
                  font=("Arial", 8), foreground="gray").pack(side="left")

        self.refresh_dropdown()

    # --- LOGIC ---

    def reload_config(self):
        """Reload the module mapping from JSON file."""
        old_count = len(self.module_map)
        self.module_map = load_module_mapping()
        new_count = len(self.module_map)
        
        # Update dropdown
        self.refresh_dropdown()
        
        messagebox.showinfo("Config Neu Geladen", 
            f"Module Mapping aktualisiert!\n\nVorher: {old_count} Module\nJetzt: {new_count} Module")

    def on_tree_select(self, event):
        """Aktiviert Buttons nur wenn etwas ausgewählt ist"""
        selected = self.tree.selection()
        state = "normal" if selected else "disabled"
        self.btn_edit.config(state=state)
        self.btn_del.config(state=state)

    def delete_selected(self, event=None):
        selected_item = self.tree.selection()
        if not selected_item: return
        
        idx = self.tree.index(selected_item)
        
        if 0 <= idx < len(self.rows_data):
            del self.rows_data[idx]
            self.recalculate_logic()
            self.refresh_ui()

    def edit_selected(self, event=None):
        selected_item = self.tree.selection()
        if not selected_item: return
        
        idx = self.tree.index(selected_item)
        if 0 <= idx < len(self.rows_data):
            row = self.rows_data[idx]
            
            self.course_var.set(row['src_key'])
            self.target_var.set(row['tgt_name'])
            self.grade_input.delete(0, tk.END)
            self.grade_input.insert(0, row['grade'])
            
            del self.rows_data[idx]
            
            self.recalculate_logic()
            self.refresh_ui()
            self.grade_input.focus_set()
            self.grade_input.select_range(0, tk.END)

    def refresh_dropdown(self):
        used = [r['src_key'] for r in self.rows_data]
        avail = sorted([k for k in self.module_map.keys() if k not in used])
        self.course_input['values'] = avail

    def on_course_select(self, event):
        selected = self.course_var.get()
        if selected in self.module_map:
            self.target_var.set(self.module_map[selected]['name'])
            self.grade_input.focus_set()

    def parse_grade(self, g):
        try: return float(g.replace(',', '.'))
        except: return 99.9

    def recalculate_logic(self):
        for r in self.rows_data: r['active'] = True
        
        # Elements Logic: Best 2
        elem_indices = []
        for i, row in enumerate(self.rows_data):
            if "Elements of" in row['src_key']:
                elem_indices.append((i, self.parse_grade(row['grade'])))
        
        if len(elem_indices) > 2:
            elem_indices.sort(key=lambda x: x[1])
            for k in range(2, len(elem_indices)):
                idx = elem_indices[k][0]
                self.rows_data[idx]['active'] = False

    def add_row_event(self, e): self.add_row()

    def add_row(self):
        key = self.course_var.get()
        grade = self.grade_input.get().strip()
        if not key or key not in self.module_map: return
        
        data = {
            'src_key': key,
            'tgt_id': self.module_map[key]['id'],
            'tgt_name': self.module_map[key]['name'],
            'grade': grade,
            'active': True
        }
        self.rows_data.append(data)
        self.recalculate_logic()
        
        self.course_input.set('')
        self.grade_input.delete(0, tk.END)
        self.target_var.set('')
        self.refresh_ui()
        self.course_input.focus_set()

    def refresh_ui(self):
        for item in self.tree.get_children(): self.tree.delete(item)
        
        for row in self.rows_data:
            stat = "Dabei" if row['active'] else "Ignoriert (Limit)"
            tags = ('normal',) if row['active'] else ('excluded',)
            self.tree.insert("", "end", values=(row['src_key'], row['grade'], row['tgt_name'], stat), tags=tags)
        
        self.refresh_dropdown()
        self.gen_btn.config(state="normal" if self.rows_data else "disabled")
        
        self.on_tree_select(None)

    def clear_all(self):
        if messagebox.askyesno("Löschen", "Alle Einträge löschen?"):
            self.rows_data = []
            self.refresh_ui()

    def simple_replace(self, paragraph, key, value):
        if key not in paragraph.text: return
        replaced = False
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, value)
                replaced = True
        if not replaced and key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)

    def generate_document(self):
        name = self.name_entry.get()
        if not name or not self.mat_entry.get():
            messagebox.showwarning("Fehler", "Bitte Name und Matrikelnummer ausfüllen.")
            return

        tpl_path = get_template_path()
        if not os.path.exists(tpl_path):
            messagebox.showerror("Fehler", "Template nicht gefunden.")
            return

        try:
            doc = Document(tpl_path)
            
            # 1. TEXT ERSETZUNG
            replacements = {
                "{Name}": name,
                "{Matrikelnummer}": self.mat_entry.get(),
                "{Matrikelnummer)": self.mat_entry.get(),
                "{Studiengang+PO}": self.prev_study_entry.get(),
                "{Studiengang}": self.prev_study_entry.get(),
                "{Fachsemester}": self.fs_entry.get(),
                "{date}": datetime.today().strftime("%d.%m.%Y"),
                "xx.xx.2025": datetime.today().strftime("%d.%m.%Y"),
                "{Herrn|Frau}": "Herrn" if self.gender_var.get() == "Herr" else "Frau",
                "{Herr/Frau}": self.gender_var.get()
            }

            def apply_replacements(p):
                for k, v in replacements.items(): self.simple_replace(p, k, v)

            for p in doc.paragraphs: apply_replacements(p)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: apply_replacements(p)
            
            # --- GRAMMATIK FIX ---
            if self.gender_var.get() == "Herr":
                wrong_start = f"Herrn {name} kann"
                correct_start = f"Herr {name} kann"
                for p in doc.paragraphs:
                    if wrong_start in p.text:
                        p.text = p.text.replace(wrong_start, correct_start)

            # 2. TABELLEN LOGIK
            target_table = None
            for t in doc.tables:
                if len(t.rows) > 0:
                    header = " ".join([c.text for c in t.rows[0].cells])
                    if "Anerkannt" in header or "Studienleistungen" in header:
                        target_table = t
                        break
            
            if target_table:
                active_entries = [d for d in self.rows_data if d['active']]
                used_entries = []

                for i in range(len(target_table.rows) - 1, 0, -1):
                    row = target_table.rows[i]
                    row_text = " ".join([c.text for c in row.cells])
                    
                    match_entry = None
                    for entry in active_entries:
                        if entry['tgt_id'] in row_text:
                            if entry in used_entries: continue 
                            match_entry = entry
                            break
                    
                    if match_entry:
                        if len(row.cells) > 0 and not row.cells[0].text.strip():
                            if "Deutsch" in match_entry['src_key']:
                                row.cells[0].text = match_entry['src_key']
                            else:
                                row.cells[0].text = match_entry['src_key'].split('(')[0]
                        
                        replaced_note = False
                        grade_val = str(match_entry['grade'])

                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if "Note" in p.text:
                                    grade_val='Note: '+grade_val
                                    p.text = p.text.replace("Note", grade_val)
                                    replaced_note = True
                        
                        if not replaced_note and len(row.cells) >= 3:
                            current_text = row.cells[-1].text.strip()
                            if current_text.endswith(":"):
                                row.cells[-1].text += f" {grade_val}"
                            else:
                                row.cells[-1].text += f": {grade_val}"

                        used_entries.append(match_entry)
                    else:
                        row._element.getparent().remove(row._element)
            
            # Save
            clean_name = name.replace(' ', '_')
            fname = f"Anerkennung_{clean_name}_{datetime.today().strftime('%Y-%m-%d')}.docx"
            out_path = os.path.join(os.path.dirname(tpl_path), fname)
            doc.save(out_path)
            
            messagebox.showinfo("Erfolg", f"Dokument erstellt:\n{out_path}")
            if sys.platform == "darwin": os.system(f"open '{out_path}'")
            elif sys.platform == "win32": os.startfile(out_path)

        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = RecognitionApp(root)
    root.mainloop()